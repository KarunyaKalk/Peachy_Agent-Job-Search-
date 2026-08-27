import os
import io
import json
import re
import httpx
from typing import Dict, Any, List, Tuple

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

from app.schemas.profile_parser import (
    ParsedResumeData,
    ParsedContactSummary,
    AmbiguityFlag,
)
from app.schemas.profile import (
    SkillCreate,
    WorkExperienceCreate,
    ExperienceBulletCreate,
    ProjectCreate,
    EducationCreate,
    CertificationCreate,
)


class ResumeParserService:
    """
    Claude API Resume Auto-Fill Engine & Document Parser.
    Extracts PDF/DOCX text and parses it into Master Profile structured schema
    using strict Fact-Guard principles and ambiguity detection.
    """

    def __init__(self):
        self.api_key = os.getenv("ANTHROPIC_API_KEY", "")
        self.model = os.getenv("LLM_MODEL", "claude-3-5-sonnet-20241022")

    def extract_text(self, file_bytes: bytes, filename: str) -> str:
        """
        Extracts plain text from PDF or DOCX file bytes.
        """
        ext = filename.lower().split(".")[-1]
        text = ""

        if ext == "pdf":
            if not pypdf:
                raise ValueError("PDF parsing library (pypdf) is not installed.")
            pdf_file = io.BytesIO(file_bytes)
            reader = pypdf.PdfReader(pdf_file)
            page_texts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    page_texts.append(t)
            text = "\n".join(page_texts)

        elif ext in ["docx", "doc"]:
            if not docx:
                raise ValueError("DOCX parsing library (python-docx) is not installed.")
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_txt = " | ".join([c.text.strip() for c in row.cells if c.text.strip()])
                    if row_txt:
                        paragraphs.append(row_txt)
            text = "\n".join(paragraphs)
        else:
            # Fallback text decoding
            try:
                text = file_bytes.decode("utf-8", errors="ignore")
            except Exception:
                raise ValueError(f"Unsupported file format: .{ext}. Please upload PDF or DOCX.")

        cleaned_text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if not cleaned_text or len(cleaned_text) < 20:
            raise ValueError("Could not extract meaningful text from the uploaded document.")

        return cleaned_text

    async def parse_resume_document(self, resume_text: str) -> Tuple[ParsedResumeData, List[AmbiguityFlag]]:
        """
        Main entry point for parsing raw resume text into structured Master Profile schema
        plus ambiguity flags using Claude API (or fallback parser if unavailable).
        """
        if self.api_key:
            try:
                parsed_data, ambiguities = await self._call_claude_parser(resume_text)
                return parsed_data, ambiguities
            except Exception as e:
                print(f"[ResumeParser Warning] Claude API call failed: {e}. Switching to fallback parser.")
                return self._fallback_resume_parser(resume_text)
        else:
            return self._fallback_resume_parser(resume_text)

    async def _call_claude_parser(self, resume_text: str) -> Tuple[ParsedResumeData, List[AmbiguityFlag]]:
        prompt = f"""
You are an expert resume parser for Peachy AI Agent.

### SOURCE RESUME TEXT:
{resume_text}

### CRITICAL INSTRUCTIONS (FACT-GUARD MODULE 3 PRINCIPLES):
1. Extract ALL details from the source document into the JSON schema below.
2. DO NOT invent, embellish, extrapolate, or hallucinate any position, employer, date, skill, metric, bullet point, or credential not explicitly present in the source text.
3. If dates are vague or incomplete (e.g., missing month, "2021 - Present" vs exact month, overlapping dates), if bullet points could belong to multiple roles, or if a section is ambiguous, DO NOT guess randomly. Add an entry to the `ambiguities` array detailing the ambiguity.

### REQUIRED JSON RESPONSE SCHEMA:
{{
  "contact": {{
    "phone": "extracted phone or null",
    "location": "extracted location (City, State) or null",
    "linkedin_url": "extracted LinkedIn URL or null",
    "github_url": "extracted GitHub URL or null",
    "portfolio_url": "extracted portfolio/website URL or null",
    "summary": "extracted summary paragraph or concise summary from profile section or null"
  }},
  "skills": [
    {{
      "category": "General | Frontend | Backend | Cloud/DevOps | Languages | Databases | Tools",
      "name": "Skill Name",
      "proficiency": "Expert | Advanced | Proficient | Optional/null"
    }}
  ],
  "experiences": [
    {{
      "company": "Company Name",
      "role": "Role Title",
      "location": "Location or null",
      "start_date": "YYYY-MM or Mon YYYY",
      "end_date": "YYYY-MM or Present or null",
      "is_current": true/false,
      "description": "Brief role overview or null",
      "bullets": [
        {{
          "content": "Verbatim or faithful bullet point text",
          "impact_category": "Metric/Quantified | Technical Achievement | Leadership | Process Improvement"
        }}
      ]
    }}
  ],
  "projects": [
    {{
      "title": "Project Title",
      "description": "Description of project",
      "technologies": "Comma-separated tech stack or null",
      "project_url": "URL or null",
      "start_date": "Start Date or null",
      "end_date": "End Date or null"
    }}
  ],
  "education": [
    {{
      "institution": "University / College Name",
      "degree": "Degree Title (e.g. Bachelor of Science in CS)",
      "field_of_study": "Field of study or null",
      "start_date": "Start Date or null",
      "end_date": "Graduation Date or null",
      "gpa_or_honors": "GPA / Honors or null"
    }}
  ],
  "certifications": [
    {{
      "name": "Certification Name",
      "issuing_organization": "Issuer Name",
      "issue_date": "Issue Date or null",
      "expiration_date": "Expiration Date or null",
      "credential_id": "ID or null",
      "credential_url": "URL or null"
    }}
  ],
  "ambiguities": [
    {{
      "id": "unique_string_id",
      "section": "contact | experience | skills | projects | education | certifications",
      "item_identifier": "Identifying string (e.g. Software Engineer at Company)",
      "field": "Field name (e.g. start_date, bullets, category)",
      "reason": "Clear explanation of why this item is ambiguous in the source text",
      "suggested_action": "Actionable recommendation for user review"
    }}
  ]
}}
"""

        async with httpx.AsyncClient() as client:
            resp = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": self.api_key,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json={
                    "model": self.model,
                    "max_tokens": 4000,
                    "temperature": 0.1,
                    "messages": [{"role": "user", "content": prompt}],
                },
                timeout=45.0,
            )

            if resp.status_code != 200:
                raise RuntimeError(f"Anthropic API returned status {resp.status_code}: {resp.text}")

            res_data = resp.json()
            content_text = res_data["content"][0]["text"]

            # Parse JSON block
            json_match = re.search(r"```(?:json)?\s*(\{[\s\S]*?\})\s*```", content_text)
            if json_match:
                raw_json = json_match.group(1)
            else:
                raw_json = content_text

            data_dict = json.loads(raw_json)

            # Build Pydantic models
            contact_dict = data_dict.get("contact", {})
            contact = ParsedContactSummary(
                phone=contact_dict.get("phone"),
                location=contact_dict.get("location"),
                linkedin_url=contact_dict.get("linkedin_url"),
                github_url=contact_dict.get("github_url"),
                portfolio_url=contact_dict.get("portfolio_url"),
                summary=contact_dict.get("summary"),
            )

            skills = [
                SkillCreate(
                    category=s.get("category", "General"),
                    name=s.get("name", "").strip(),
                    proficiency=s.get("proficiency"),
                )
                for s in data_dict.get("skills", [])
                if s.get("name")
            ]

            experiences = []
            for exp in data_dict.get("experiences", []):
                bullets = [
                    ExperienceBulletCreate(
                        content=b.get("content", "").strip(),
                        impact_category=b.get("impact_category"),
                    )
                    for b in exp.get("bullets", [])
                    if b.get("content")
                ]
                experiences.append(
                    WorkExperienceCreate(
                        company=exp.get("company", "Company"),
                        role=exp.get("role", "Role"),
                        location=exp.get("location"),
                        start_date=exp.get("start_date", "2020"),
                        end_date=exp.get("end_date"),
                        is_current=bool(exp.get("is_current", False)),
                        description=exp.get("description"),
                        bullets=bullets,
                    )
                )

            projects = [
                ProjectCreate(
                    title=p.get("title", "Project"),
                    description=p.get("description"),
                    technologies=p.get("technologies"),
                    project_url=p.get("project_url"),
                    start_date=p.get("start_date"),
                    end_date=p.get("end_date"),
                )
                for p in data_dict.get("projects", [])
                if p.get("title")
            ]

            education = [
                EducationCreate(
                    institution=e.get("institution", "Institution"),
                    degree=e.get("degree", "Degree"),
                    field_of_study=e.get("field_of_study"),
                    start_date=e.get("start_date"),
                    end_date=e.get("end_date"),
                    gpa_or_honors=e.get("gpa_or_honors"),
                )
                for e in data_dict.get("education", [])
                if e.get("institution")
            ]

            certifications = [
                CertificationCreate(
                    name=c.get("name", "Certification"),
                    issuing_organization=c.get("issuing_organization", "Organization"),
                    issue_date=c.get("issue_date"),
                    expiration_date=c.get("expiration_date"),
                    credential_id=c.get("credential_id"),
                    credential_url=c.get("credential_url"),
                )
                for c in data_dict.get("certifications", [])
                if c.get("name")
            ]

            ambiguities = [
                AmbiguityFlag(
                    id=a.get("id", f"amb_{idx}"),
                    section=a.get("section", "experience"),
                    item_identifier=a.get("item_identifier"),
                    field=a.get("field", "dates"),
                    reason=a.get("reason", "Ambiguous detail in resume source text."),
                    suggested_action=a.get("suggested_action", "Please review and adjust manually."),
                )
                for idx, a in enumerate(data_dict.get("ambiguities", []))
            ]

            parsed_data = ParsedResumeData(
                contact=contact,
                summary=contact.summary,
                skills=skills,
                experiences=experiences,
                projects=projects,
                education=education,
                certifications=certifications,
            )

            return parsed_data, ambiguities

    def _fallback_resume_parser(self, resume_text: str) -> Tuple[ParsedResumeData, List[AmbiguityFlag]]:
        """
        Regex & heuristic fallback parser for when Claude API is not configured or fails.
        Extracts contact info, summary, experience, skills, projects, education, certs.
        """
        ambiguities: List[AmbiguityFlag] = []

        # 1. Contact Info Extraction
        phone_match = re.search(r"(\+?\d{1,3}[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}", resume_text)
        email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", resume_text)
        linkedin_match = re.search(r"https?://(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+", resume_text, re.I)
        github_match = re.search(r"https?://(?:www\.)?github\.com/[a-zA-Z0-9_-]+", resume_text, re.I)
        portfolio_match = re.search(r"https?://(?:www\.)?[a-zA-Z0-9_-]+\.(?:io|com|dev|me|net)", resume_text, re.I)

        # Location heuristic
        location_match = re.search(r"([A-Z][a-zA-B\s]+,\s*(?:[A-Z]{2}|[A-Z][a-z]+))", resume_text)

        contact = ParsedContactSummary(
            phone=phone_match.group(0) if phone_match else None,
            location=location_match.group(1) if location_match else None,
            linkedin_url=linkedin_match.group(0) if linkedin_match else None,
            github_url=github_match.group(0) if github_match else None,
            portfolio_url=portfolio_match.group(0) if portfolio_match else None,
        )

        lines = [line.strip() for line in resume_text.split("\n") if line.strip()]

        # Section segmentation
        sections: Dict[str, List[str]] = {
            "summary": [],
            "experience": [],
            "skills": [],
            "projects": [],
            "education": [],
            "certifications": [],
        }

        current_sec = "summary"
        for line in lines:
            lower = line.lower()
            if any(k in lower for k in ["work experience", "professional experience", "employment history", "experience"]):
                current_sec = "experience"
                continue
            elif any(k in lower for k in ["skills", "technical skills", "technologies", "expertise"]):
                current_sec = "skills"
                continue
            elif any(k in lower for k in ["projects", "personal projects", "key projects"]):
                current_sec = "projects"
                continue
            elif any(k in lower for k in ["education", "academic background", "degrees"]):
                current_sec = "education"
                continue
            elif any(k in lower for k in ["certifications", "licenses", "certificates"]):
                current_sec = "certifications"
                continue
            elif any(k in lower for k in ["summary", "about me", "profile summary", "objective"]):
                current_sec = "summary"
                continue

            sections[current_sec].append(line)

        # Summary text
        summary_text = " ".join(sections["summary"][:4]) if sections["summary"] else None
        contact.summary = summary_text

        # Skills extraction
        skills: List[SkillCreate] = []
        raw_skills = sections["skills"]
        for s_line in raw_skills:
            parts = re.split(r"[:,•·|]", s_line)
            cat = "General"
            if len(parts) > 1 and len(parts[0].split()) < 4:
                cat = parts[0].strip()
                items = parts[1:]
            else:
                items = parts

            for item in items:
                for sub in item.split(","):
                    name = sub.strip(" •·,;:")
                    if name and len(name) < 40 and not name.lower().startswith("skills"):
                        skills.append(SkillCreate(category=cat, name=name))

        # Experience extraction
        experiences: List[WorkExperienceCreate] = []
        exp_lines = sections["experience"]
        current_exp: WorkExperienceCreate = None

        date_pattern = r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|[0-9]{4})[\s\S]*?(?:Present|[0-9]{4}))"

        for line in exp_lines:
            dates = re.findall(date_pattern, line, re.I)
            is_bullet = line.startswith("-") or line.startswith("•") or line.startswith("*") or line.startswith("·")

            if (dates or (len(line.split()) < 8 and not is_bullet)) and not current_exp:
                # New role heading candidate
                current_exp = WorkExperienceCreate(
                    company="Extracted Company",
                    role=line[:60],
                    start_date=dates[0] if dates else "2021",
                    end_date="Present" if "present" in line.lower() else None,
                    is_current="present" in line.lower(),
                    bullets=[],
                )
                experiences.append(current_exp)
            elif is_bullet and current_exp:
                bullet_content = line.lstrip("-•*· ").strip()
                if bullet_content:
                    current_exp.bullets.append(ExperienceBulletCreate(content=bullet_content))
            elif current_exp and len(line) > 10:
                if not current_exp.description:
                    current_exp.description = line
                else:
                    current_exp.bullets.append(ExperienceBulletCreate(content=line))

        # Add ambiguity check for extracted experience
        if experiences:
            for exp in experiences:
                if exp.start_date == "2021":
                    ambiguities.append(
                        AmbiguityFlag(
                            id=f"amb_date_{exp.company}",
                            section="experience",
                            item_identifier=f"{exp.role} at {exp.company}",
                            field="start_date",
                            reason="Date range was extracted using fallback parser pattern matching.",
                            suggested_action="Verify exact start and end employment dates.",
                        )
                    )

        # Education extraction
        education: List[EducationCreate] = []
        for edu_line in sections["education"]:
            if len(edu_line) > 5:
                education.append(
                    EducationCreate(
                        institution=edu_line.split(",")[0] if "," in edu_line else edu_line,
                        degree="Degree / Studies",
                        field_of_study="Computer Science & Engineering",
                    )
                )

        # Projects extraction
        projects: List[ProjectCreate] = []
        for proj_line in sections["projects"]:
            if len(proj_line) > 5:
                projects.append(ProjectCreate(title=proj_line[:50], description=proj_line))

        # Certifications extraction
        certifications: List[CertificationCreate] = []
        for cert_line in sections["certifications"]:
            if len(cert_line) > 3:
                certifications.append(
                    CertificationCreate(
                        name=cert_line[:60],
                        issuing_organization="Issuing Body",
                    )
                )

        parsed_data = ParsedResumeData(
            contact=contact,
            summary=summary_text,
            skills=skills[:25],
            experiences=experiences[:10],
            projects=projects[:5],
            education=education[:3],
            certifications=certifications[:5],
        )

        return parsed_data, ambiguities
